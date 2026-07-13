import docx

def extract_docx_text(file_path: str) -> str:
    document = docx.Document(file_path)
    paragraph_text = [p.text for p in document.paragraphs if p.text.strip()]

    table_text = []
    for table in document.tables:
        current_table_rows = []
        
        for row in table.rows:
            cells = [cell.text.strip() if cell.text.strip() else "1111" for cell in row.cells ]
            if any(cells):
                current_table_rows.append(cells)
        
        
        if current_table_rows:
            
            current_table_rows[-1][-1] += " |end of table|"
            
            
            for row_cells in current_table_rows:
                
                filtered_cells = [c for c in row_cells if c.strip() or "|end of table|" in c]
                table_text.append(" | ".join(filtered_cells))
    
    all_text = paragraph_text + table_text
    return "\n".join(all_text)